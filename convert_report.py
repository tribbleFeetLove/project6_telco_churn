"""将实验报告.md转换为HTML和DOCX格式"""
import re, os

BASE = os.path.dirname(os.path.abspath(__file__))
md_path = os.path.join(BASE, '实验报告.md')

with open(md_path, 'r', encoding='utf-8') as f:
    md = f.read()

# ===== Generate HTML =====
md_html = md

# Headers
md_html = re.sub(r'^#### (.*)', r'<h4>\1</h4>', md_html, flags=re.M)
md_html = re.sub(r'^### (.*)', r'<h3>\1</h3>', md_html, flags=re.M)
md_html = re.sub(r'^## (.*)', r'<h2>\1</h2>', md_html, flags=re.M)
md_html = re.sub(r'^# (.*)', r'<h1>\1</h1>', md_html, flags=re.M)

# Bold / italic
md_html = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', md_html)
md_html = re.sub(r'\*(.*?)\*', r'<i>\1</i>', md_html)

# Inline code
md_html = re.sub(r'`([^`]+)`', r'<code>\1</code>', md_html)

# Table handling
lines = md_html.split('\n')
out_lines = []
table_buf = []
in_table = False
for line in lines:
    stripped = line.strip()
    if stripped.startswith('|') and stripped.endswith('|'):
        if not in_table:
            in_table = True
            table_buf = []
        if not re.match(r'^\|[\s\-:|]+\|$', stripped):
            table_buf.append(stripped)
    else:
        if in_table:
            html = '<table border="1" cellpadding="4" cellspacing="0" style="border-collapse:collapse;margin:10px 0;">\n'
            for j, tl in enumerate(table_buf):
                cells = [c.strip() for c in tl.split('|')[1:-1]]
                tag = 'th' if j == 0 else 'td'
                html += '<tr>' + ''.join(f'<{tag} style="padding:4px 8px;">{c}</{tag}>' for c in cells) + '</tr>\n'
            html += '</table>'
            out_lines.append(html)
            in_table = False
            table_buf = []
        out_lines.append(line)

md_html = '\n'.join(out_lines)

# Code blocks
def replace_codeblock(m):
    return '<pre><code>' + m.group(1) + '</code></pre>'
md_html = re.sub(r'```(.*?)```', replace_codeblock, md_html, flags=re.S)

# Lists
md_html = re.sub(r'^- (.*)', r'<li>\1</li>', md_html, flags=re.M)

html_content = f'''<!DOCTYPE html>
<html><head><meta charset="UTF-8">
<style>
body {{ font-family: '宋体', SimSun, serif; font-size: 12pt; line-height: 1.8; max-width: 210mm; margin: 20px auto; padding: 0 20px; color: #000; }}
h1 {{ font-size: 18pt; text-align: center; font-weight: bold; margin: 20px 0; }}
h2 {{ font-size: 15pt; font-weight: bold; margin: 16px 0 10px 0; border-bottom: 1px solid #333; padding-bottom: 5px; }}
h3 {{ font-size: 13pt; font-weight: bold; margin: 12px 0 8px 0; }}
h4 {{ font-size: 12pt; font-weight: bold; margin: 10px 0 6px 0; }}
table {{ border-collapse: collapse; width: 100%; margin: 10px 0; font-size: 10pt; }}
th {{ background-color: #4472C4; color: white; font-weight: bold; text-align: center; }}
td {{ text-align: center; }}
tr:nth-child(even) {{ background-color: #f2f2f2; }}
code {{ background: #f5f5f5; padding: 1px 3px; font-family: Consolas, monospace; font-size: 10pt; }}
pre {{ background: #f5f5f5; padding: 10px; font-size: 9pt; overflow-x: auto; }}
p {{ text-indent: 2em; margin: 6px 0; }}
ul {{ margin-left: 2em; }}
</style></head><body>
{md_html}
</body></html>'''

with open(os.path.join(BASE, '实验报告.html'), 'w', encoding='utf-8') as f:
    f.write(html_content)
print('HTML 生成成功')

# ===== Generate DOCX =====
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Default style
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)

# Title
title = doc.add_heading('电信客户流失预测与价值细分系统', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.add_run('《数据分析与数据挖掘》课程项目报告 — 项目6').font.size = Pt(14)

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run('小组成员：__________\n专业班级：__________\n指导老师：__________')

doc.add_page_break()

# Process sections
# Split by ## headers
sections = re.split(r'\n## ', '\n' + md)
for si, section in enumerate(sections):
    if si == 0:
        # Before first ## : abstract etc
        subs = re.split(r'\n### ', section)
        for sub_text in subs:
            lines = sub_text.strip().split('\n', 1)
            if len(lines) >= 1 and not lines[0].startswith('#'):
                hdr = lines[0].strip()
                body = lines[1].strip() if len(lines) > 1 else ''
                if hdr == '摘要' or hdr == '关键词':
                    doc.add_heading(hdr, level=1)
                if body:
                    for para in body.split('\n\n'):
                        para = para.strip()
                        if para:
                            doc.add_paragraph(para)
            elif lines[0].startswith('# '):
                doc.add_heading(lines[0].replace('# ', ''), level=1)
                if len(lines) > 1:
                    doc.add_paragraph(lines[1].strip())
        continue

    # Parse section header and body
    parts = section.split('\n', 1)
    sec_title = parts[0].strip().lstrip('#').strip()
    sec_body = parts[1] if len(parts) > 1 else ''

    # Add heading (skip TOC numbers like "1. ")
    clean_title = re.sub(r'^\d+\.\s*', '', sec_title)
    if clean_title:
        doc.add_heading(clean_title, level=1)

    # Sub-sections
    subs = re.split(r'\n### ', sec_body)
    for sub_text in subs:
        sub_parts = sub_text.split('\n', 1)
        sub_hdr = sub_parts[0].strip()
        sub_content = sub_parts[1] if len(sub_parts) > 1 else ''

        clean_sub = re.sub(r'^\d+\.\d+\s*', '', sub_hdr)
        if clean_sub and not clean_sub.startswith('#'):
            doc.add_heading(clean_sub, level=2)

        # Paragraphs
        for para in sub_content.split('\n\n'):
            para = para.strip()
            if not para:
                continue
            # Skip table markdown
            if para.startswith('|'):
                continue
            # Clean formatting
            clean = re.sub(r'\*\*(.*?)\*\*', r'\1', para)
            clean = re.sub(r'\*(.*?)\*', r'\1', clean)
            clean = re.sub(r'`([^`]+)`', r'\1', clean)
            clean = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', clean)
            clean = re.sub(r'!\[.*?\]\(.*?\)', '', clean)
            doc.add_paragraph(clean)

# Add note at the end
doc.add_paragraph('\n注：完整的图表请参见 experiments/ 目录。本报告可用浏览器打开 实验报告.html 获得更好的表格和公式渲染效果。')

doc_path = os.path.join(BASE, '实验报告.docx')
doc.save(doc_path)
print(f'DOCX 生成成功: {doc_path}')
