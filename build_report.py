"""生成格式规范的实验报告 DOCX"""
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os, re

BASE = r'C:\Users\asus\Desktop\数据挖掘\project6_telco_churn'
doc = Document()

# ====== 全局样式设置 ======
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
pf = style.paragraph_format
pf.line_spacing = 1.5
pf.space_after = Pt(6)
pf.first_line_indent = Cm(0.74)  # 两字符缩进

# 标题样式
for lvl in range(1, 4):
    hs = doc.styles[f'Heading {lvl}']
    hs.font.name = '黑体'
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    hs.font.color.rgb = RGBColor(0, 0, 0)
    hs.font.bold = True
    hs.paragraph_format.first_line_indent = Cm(0)
    if lvl == 1:
        hs.font.size = Pt(16)
    elif lvl == 2:
        hs.font.size = Pt(14)
    else:
        hs.font.size = Pt(12)

def add_para(text, bold=False, indent=True, size=12, align=None, font_name=None):
    """添加段落"""
    p = doc.add_paragraph()
    if not indent:
        p.paragraph_format.first_line_indent = Cm(0)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    if font_name:
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    return p

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.first_line_indent = Cm(0)
    return h

def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    run.font.size = Pt(12)
    return p

def add_table(headers, rows, col_widths=None):
    """添加格式化表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        # 蓝色背景
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="4472C4" w:val="clear"/>')
        cell._tc.get_or_add_tcPr().append(shading)
        run.font.color.rgb = RGBColor(255, 255, 255)
    # 数据行
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i+1].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            # 交替行背景
            if i % 2 == 1:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9E2F3" w:val="clear"/>')
                cell._tc.get_or_add_tcPr().append(shading)
    # 列宽
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()  # 表后空行
    return table

def add_formula(text):
    """添加公式（用等宽字体+缩进表示）"""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(11)
    run.italic = True
    return p

# ===================================================================
# 封面
# ===================================================================
for _ in range(6):
    doc.add_paragraph()

add_para('《数据分析与数据挖掘》课程项目报告', bold=True, size=22, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
doc.add_paragraph()
add_para('电信客户流失预测与价值细分系统', bold=True, size=20, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
add_para('—— 项目6：基于机器学习的客户流失分析与业务策略', size=14, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)

for _ in range(4):
    doc.add_paragraph()

cover_info = [
    ('小组成员', '__________  __________  __________'),
    ('专业班级', '__________'),
    ('指导老师', '__________'),
    ('开课时间', '2025-2026-1'),
]
for label, value in cover_info:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(f'{label}：{value}')
    run.font.size = Pt(14)

doc.add_page_break()

# ===================================================================
# 摘要
# ===================================================================
add_heading('摘要', level=1)
add_para('电信行业客户流失问题日益严峻，新客户获取成本远高于存量客户维护成本。本项目构建了一套完整的"流失预警+客户分群"双模块分析系统，基于Kaggle Telco Customer Churn数据集（7,043条记录，21个特征），综合运用数据清洗、RFM与WOE特征工程、SMOTE过采样、四种分类模型对比（逻辑回归、随机森林、XGBoost、LightGBM）、SHAP可解释性分析以及K-Means/DBSCAN聚类分析等方法。')
add_para('实验结果表明，在真实不平衡测试集上，随机森林模型取得最佳综合性能（ROC-AUC = 0.8382，Recall = 0.7139，F1 = 0.6350）。K-Means聚类成功将客户分为4个具有显著差异的群体，其中高危群体（Cluster 2）流失率高达60.9%，高价值群体（Cluster 1）月均消费$95且流失率仅21.3%。基于SHAP分析和聚类画像，提出了合同类型优化、增值服务捆绑、支付方式引导等六大方向的业务策略建议，为电信运营商的客户留存决策提供了数据驱动的支撑。')
add_para('关键词：客户流失预测；机器学习；SMOTE；SHAP可解释性；客户细分；集成学习', bold=True, indent=False)

# ===================================================================
# 1. 引言
# ===================================================================
add_heading('1. 引言', level=1)

add_heading('1.1 问题定义', level=2)
add_para('在电信行业，客户流失（Customer Churn）指客户主动停止使用运营商服务的行为。随着市场竞争日趋激烈，运营商之间的客户争夺成本不断攀升。研究表明，获取一个新客户的成本是维护一个老客户的5至6倍，而客户留存率每提升5%，企业利润可增加25%至95%[1]。因此，准确预测客户的流失倾向、深入理解流失驱动因素、制定针对性的客户留存策略，对于电信企业保持竞争力和盈利能力具有重要的商业价值。')

add_heading('1.2 研究意义', level=2)
add_para('传统的客户流失管理往往采取被动的事后补救措施，效果有限且成本高昂。数据挖掘与机器学习技术的快速发展为主动式流失管理提供了技术基础。通过构建预测模型，企业可以在客户做出离网决策之前识别高风险客户群体，从而提前采取精准的干预措施。此外，借助SHAP（SHapley Additive exPlanations）等可解释性工具，模型分析结果可以由"黑箱预测"转变为可理解、可操作的业务洞察[2]。')

add_heading('1.3 项目目标', level=2)
add_para('本项目的目标是构建一个完整的"流失预警 + 客户分群"双模块分析系统。具体包括三个层次：')
add_bullet('有监督学习模块：对比逻辑回归、随机森林、XGBoost和LightGBM四种分类算法，通过SMOTE处理类别不平衡，以ROC-AUC和Recall为核心指标，选出最佳流失预测模型。')
add_bullet('无监督学习模块：利用K-Means和DBSCAN对客户进行聚类细分，识别不同价值群体及其流失风险特征。')
add_bullet('可解释性与业务应用：通过SHAP分析揭示特征对流失预测的贡献方向与大小，结合聚类画像，输出可落地的业务策略建议。')

# ===================================================================
# 2. 相关工作
# ===================================================================
add_heading('2. 相关工作', level=1)
add_para('客户流失预测是数据挖掘领域的经典问题，国内外学者已开展了广泛的研究工作。本节从传统统计方法、集成学习、不平衡学习策略、深度学习方法以及模型可解释性五个方面进行综述。')

add_heading('2.1 传统统计方法', level=2)
add_para('早期研究主要采用逻辑回归（Logistic Regression）和决策树（Decision Tree）等传统方法。Lemmens和Croux[3]使用逻辑回归对美国电信公司客户数据建模，发现合同类型和在网时长是最具预测力的变量。该类方法的优势在于模型可解释性强，训练和推理速度快，但面对高维特征和非线性关系时预测能力有限。Burez和Van den Poel[6]的研究进一步指出，单一的传统分类器在处理类别不平衡时性能下降明显。')

add_heading('2.2 集成学习方法', level=2)
add_para('随着算法发展，集成学习方法在流失预测中展现出优越性能。XGBoost（eXtreme Gradient Boosting）由Chen和Guestrin[4]于2016年提出，通过梯度提升框架实现高效的特征选择和L1/L2正则化，已成为工业界的主流方案。LightGBM由Ke等人[5]于2017年提出，基于直方图算法和叶子优先（leaf-wise）的生长策略，进一步优化了训练速度与内存占用，在大规模数据集上表现突出。Breiman[13]提出的随机森林则通过Bagging策略和多棵决策树的投票机制，天然具有良好的抗过拟合能力。')

add_heading('2.3 不平衡学习策略', level=2)
add_para('客户流失数据通常存在严重的类别不平衡问题（流失客户占比往往不到30%）。Chawla等人[7]于2002年提出的SMOTE（Synthetic Minority Oversampling Technique）通过在特征空间中合成少数类样本，有效缓解了不平衡问题，成为该领域引用最高的方法之一。Verbeke等人[8]从利润驱动的角度出发，提出代价敏感学习（Cost-Sensitive Learning）框架，通过为少数类赋予更高的误分类代价来优化模型的实际商业价值。')

add_heading('2.4 深度学习与图神经网络方法', level=2)
add_para('近年来，深度学习也被引入客户流失预测领域。Wangperawong等人[9]使用时间卷积网络（TCN）建模客户行为时序特征，取得了优于传统方法的预测精度。然而，Grinsztajn等人[10]在2022年的系统性研究中指出，对于中小规模的结构化表格数据，精心调参的集成树模型（如XGBoost、LightGBM）往往不逊色于甚至优于深度神经网络。')

add_heading('2.5 模型可解释性', level=2)
add_para('随着机器学习模型在业务场景中的广泛应用，模型可解释性变得愈发重要。Ribeiro等人[11]于2016年提出了LIME（Local Interpretable Model-agnostic Explanations），通过局部线性近似来解释单个预测。Lundberg和Lee[2]于2017年提出的SHAP基于博弈论中的Shapley值，为每个特征分配对预测结果的贡献值，兼具局部解释和全局解释的能力。在企业实际应用中，SHAP已成为解释树模型预测的主流工具。')

add_para('本项目在上述研究基础上，综合运用SMOTE过采样、四种分类模型对比、SHAP解释和聚类分析，构建一个完整、可解释、可落地的客户流失预测与价值细分系统。')

# ===================================================================
# 3. 数据集
# ===================================================================
add_heading('3. 数据集', level=1)

add_heading('3.1 数据来源与规模', level=2)
add_para('本实验使用Kaggle平台公开的Telco Customer Churn Dataset[12]。该数据集模拟了某电信运营商7,043名客户的订阅信息，包含21个特征变量（含1个目标变量），被广泛应用于客户流失预测的教学与研究。')

add_para('数据集的特征可分为三大类别，如表1所示。', indent=False)

add_table(
    ['类别', '特征名称', '说明'],
    [
        ['人口统计', 'gender, SeniorCitizen, Partner, Dependents', '性别、是否老年人、是否有伴侣、是否有家属'],
        ['账户信息', 'tenure, Contract, PaperlessBilling, PaymentMethod, MonthlyCharges, TotalCharges', '在网时长、合同类型、无纸化账单、支付方式、月费、总费'],
        ['服务订阅', 'PhoneService, MultipleLines, InternetService, OnlineSecurity, OnlineBackup, DeviceProtection, TechSupport, StreamingTV, StreamingMovies', '电话服务、多线路、互联网服务类型、在线安全等9项增值服务'],
        ['目标变量', 'Churn', '是否流失（Yes/No）'],
    ]
)

add_heading('3.2 数据分布分析', level=2)
add_para('目标变量Churn的分布存在明显的类别不平衡：在7,043名客户中，5,174名未流失（73.46%），1,869名已流失（26.54%），流失与非流失比例约为1:2.8。图1展示了目标变量的分布情况（见附录experiments/target_distribution.png）。')

add_heading('3.3 数据预处理流程', level=2)
add_para('数据预处理包含以下五个步骤：')
add_bullet('缺失值处理：经检查，数据集无显式NaN缺失值，但TotalCharges列存在11个空字符串记录（对应tenure = 0的新客户），将其转为NaN后用中位数填充。')
add_bullet('无关特征移除：移除customerID列，该列为客户唯一标识符，对预测任务无贡献。')
add_bullet('目标变量编码：将Churn列中的"Yes"/"No"映射为1/0，便于模型处理。')
add_bullet('类别特征编码：对15个类别型特征（gender、Partner、Contract等）采用Label Encoding进行数值化编码。')
add_bullet('异常值检测：采用IQR（四分位距）方法对tenure、MonthlyCharges、TotalCharges三个核心数值特征进行检测。检测结果显示数据质量较好，无超出1.5倍IQR范围的异常值。')

add_heading('3.4 数据增强策略', level=2)
add_para('针对类别不平衡问题（流失:非流失 ≈ 1:2.8），采用SMOTE进行数据增强。SMOTE通过在特征空间中为每个少数类（流失客户）样本选择k个近邻，在样本与其近邻之间随机线性插值生成合成样本。SMOTE处理后，训练集样本数由5,634增至8,278，两类样本达到1:1平衡。值得注意的是，SMOTE仅在训练集上执行，测试集保持原始分布以真实评估模型泛化能力。')

# ===================================================================
# 4. 方法
# ===================================================================
add_heading('4. 方法', level=1)

add_heading('4.1 系统总体架构', level=2)
add_para('本项目的技术架构分为四个阶段：数据层（数据加载、清洗、EDA探索性分析）、特征层（RFM特征构建、WOE分箱编码、特征重要性预筛选）、模型层（SMOTE过采样、四种分类模型训练与对比、K-Means/DBSCAN客户聚类）、应用层（SHAP可解释性分析、聚类画像、业务策略输出）。')

add_heading('4.2 特征工程', level=2)
add_para('特征工程是本项目的核心环节，包括RFM特征构建和WOE分箱编码两部分。', indent=False)

add_para('（1）RFM特征构建', bold=True, indent=False)
add_para('借鉴零售领域的RFM（Recency-Frequency-Monetary）模型，为电信场景定制三个复合特征。R指标（Recency）定义为 R = 1/(tenure + 1)，数值越高表示客户越"新"，处于流失高风险期。F指标（Frequency）统计客户订阅的增值服务数量（共9项），反映使用频率和客户粘性。M指标（Monetary）直接用MonthlyCharges衡量客户消费价值。此外，构建RFM综合评分以捕捉三维特征的联合效应。')

add_para('（2）WOE分箱编码', bold=True, indent=False)
add_para('Weight of Evidence（证据权重）是一种有监督的离散化编码方法，尤其适用于逻辑回归等线性模型。WOE的计算方式为：对每个分箱，计算好客户（未流失）与坏客户（流失）在该分箱中占比之比的对数值，即 WOE = ln[(好客户占比) / (坏客户占比)]。WOE值反映了每个分箱中两类客户的比例差异，绝对值越大表示该分箱对预测目标的区分能力越强。本实验对tenure、MonthlyCharges、TotalCharges三个核心数值特征采用等频分箱（q = 10），生成对应的WOE编码特征。')

add_heading('4.3 分类模型', level=2)
add_para('本实验对比四种分类算法，参数设置如表2所示。', indent=False)

add_table(
    ['模型', '类型', '关键参数', '特点'],
    [
        ['逻辑回归', '线性分类器', 'C=1.0, solver=liblinear, max_iter=1000', '可解释性强，训练快速，作为线性基线'],
        ['随机森林', 'Bagging集成', 'n_estimators=200, max_depth=10, min_samples_split=10', '抗过拟合，特征重要性评估，无需特征缩放'],
        ['XGBoost', 'Boosting集成', 'n_estimators=200, max_depth=6, lr=0.1, subsample=0.8', 'L1/L2正则化，梯度提升框架，工业界首选'],
        ['LightGBM', 'Boosting集成', 'n_estimators=200, max_depth=6, lr=0.1, num_leaves=31', '直方图算法，leaf-wise生长，速度快，内存省'],
    ]
)

add_heading('4.4 损失函数与优化器', level=2)
add_para('逻辑回归使用交叉熵损失（Cross-Entropy Loss），优化器为liblinear（坐标下降法）。XGBoost和LightGBM使用二分类对数损失（Binary Logloss），优化器为各自的梯度提升框架。随机森林基于基尼不纯度（Gini Impurity）进行节点分裂，不需要梯度优化。')

add_heading('4.5 评估策略', level=2)
add_para('评估指标包括Accuracy（准确率）、Precision（精确率）、Recall（召回率）、F1-Score、ROC-AUC和PR-AUC六项。其中ROC-AUC是模型区分能力的综合指标，Recall直接反映流失客户的识别率，PR-AUC在不平衡场景下比ROC-AUC更具参考价值。验证方法采用5折分层交叉验证（StratifiedKFold, shuffle = True），数据按80%训练/20%测试分层划分（stratify = y），SMOTE仅在训练集上执行。随机种子固定为42，确保实验完全可复现。')

add_heading('4.6 客户聚类', level=2)
add_para('K-Means聚类基于RFM特征（tenure, MonthlyCharges, TotalCharges, R_Score, F_Score, M_Score, RFM_Total），先通过标准化消除量纲差异，再使用肘部法则（Elbow Method）和轮廓系数（Silhouette Score）联合确定最佳K值（K = 4）。为对比验证，同时采用DBSCAN（eps = 0.5, min_samples = 10）进行基于密度的聚类，识别异常客户群体（噪声点）。')

add_heading('4.7 模型可解释性', level=2)
add_para('使用SHAP（SHapley Additive exPlanations）进行模型解释。SHAP值基于博弈论中Shapley值的理论框架，将每个预测结果分解为各特征的贡献之和。对于树模型，采用TreeExplainer实现高效的SHAP值计算；分析时采样500个训练样本以平衡计算效率与统计代表性。')

# ===================================================================
# 5. 实验
# ===================================================================
add_heading('5. 实验', level=1)

add_heading('5.1 实验设计', level=2)
add_para('本实验包含三个层次的对比：（1）模型层面——逻辑回归 vs 随机森林 vs XGBoost vs LightGBM；（2）特征层面——原始特征（19维）vs 原始特征 + RFM + WOE编码（27维）的消融对比；（3）聚类层面——K-Means（固定簇数）vs DBSCAN（密度自适应）。')

add_heading('5.2 5折交叉验证结果', level=2)
add_para('在SMOTE处理后的平衡训练集（8,278样本）上进行5折分层交叉验证，结果如表3所示。', indent=False)

add_table(
    ['模型', 'Accuracy', 'Precision', 'Recall', 'F1', 'ROC-AUC'],
    [
        ['Logistic Regression', '0.8195 ± 0.0076', '0.7977 ± 0.0076', '0.8562 ± 0.0136', '0.8259 ± 0.0079', '0.9089 ± 0.0039'],
        ['Random Forest', '0.8384 ± 0.0099', '0.8193 ± 0.0093', '0.8683 ± 0.0150', '0.8430 ± 0.0101', '0.9175 ± 0.0035'],
        ['XGBoost', '0.8530 ± 0.0089', '0.8531 ± 0.0077', '0.8529 ± 0.0140', '0.8529 ± 0.0094', '0.9337 ± 0.0040'],
        ['LightGBM', '0.8524 ± 0.0071', '0.8499 ± 0.0068', '0.8560 ± 0.0097', '0.8529 ± 0.0073', '0.9325 ± 0.0041'],
    ]
)

add_para('在平衡后的训练集上，四种模型均表现良好，XGBoost以ROC-AUC = 0.9337位居第一，体现了梯度提升集成学习的优势。四种模型的标准差均较小（< 0.015），表明模型性能稳定可靠。')

add_heading('5.3 测试集评估结果', level=2)
add_para('在保持原始分布（流失率26.54%）的真实测试集上评估，结果如表4所示。这是最能反映模型实际应用价值的评估场景。', indent=False)

add_table(
    ['模型', 'Accuracy', 'Precision', 'Recall', 'F1', 'ROC-AUC', 'PR-AUC'],
    [
        ['Logistic Regression', '0.7594', '0.5362', '0.6925', '0.6044', '0.8359', '0.6587'],
        ['Random Forest', '0.7821', '0.5717', '0.7139', '0.6350', '0.8382', '0.6527'],
        ['XGBoost', '0.7779', '0.5788', '0.5989', '0.5887', '0.8220', '0.6313'],
        ['LightGBM', '0.7764', '0.5774', '0.5882', '0.5828', '0.8220', '0.6353'],
    ]
)

add_para('在真实不平衡数据上，随机森林以ROC-AUC = 0.8382和Recall = 0.7139取得综合最佳表现。这意味着模型能够成功识别71.4%的流失客户，同时保持57.2%的精确率。逻辑回归虽然整体准确率最低，但其Recall（流失客户召回率）达到0.6925，与集成方法相当，说明线性模型在RFM和WOE特征工程的加持下仍具有竞争力。XGBoost和LightGBM在平衡训练集上表现突出（ROC-AUC > 0.93），但在真实不平衡测试集上存在一定程度的过拟合。')

add_heading('5.4 ROC与PR曲线分析', level=2)
add_para('四种模型的ROC曲线均显著优于随机猜测基线（AUC = 0.50）。（完整ROC和PR曲线图见附录experiments/roc_curves.png和pr_curves.png）。在不平衡分类问题中，PR曲线比ROC曲线更具参考价值。由于流失客户仅占26.5%，随机猜测的PR-AUC基线约为0.27。四种模型的PR-AUC在0.63至0.66之间，远超基线，说明模型具有良好的实用价值。')

add_heading('5.5 混淆矩阵分析', level=2)
add_para('从混淆矩阵可以看出（见附录experiments/confusion_matrices.png），四种模型对非流失客户的识别准确率较高（真负例多），但对流失客户存在一定程度的漏检（假负例）。以随机森林为例，374个真实流失客户中有267个被正确识别（Recall = 71.4%），107个被漏检。这是不平衡分类的典型特征——模型倾向于预测多数类。在实际业务中，可通过调整分类阈值来平衡精确率与召回率。')

add_heading('5.6 特征重要性分析', level=2)
add_para('随机森林模型的特征重要性排序显示（见附录experiments/feature_importance.png），Contract（合同类型）是最重要的预测特征（重要性 = 0.139），其次是R_Score（重要性 = 0.094）、tenure（重要性 = 0.082）和MonthlyCharges（重要性 = 0.075）。RFM综合评分RFM_Total和WOE编码特征也进入了Top-10。')

add_heading('5.7 消融实验', level=2)
add_para('为验证RFM + WOE特征工程的有效性，使用仅含原始特征（无RFM/WOE，19维）的数据集进行消融对比，结果如表5所示。', indent=False)

add_table(
    ['实验条件', 'Logistic Regression AUC', 'Random Forest AUC'],
    [
        ['原始特征（19维）', '0.8035', '0.8221'],
        ['原始特征 + RFM + WOE（27维）', '0.8359 (+0.032)', '0.8382 (+0.016)'],
    ]
)

add_para('特征工程为逻辑回归带来了+0.032的显著AUC提升，为随机森林带来+0.016的提升。这说明RFM和WOE特征能够为线性模型提供额外的非线性区分能力，验证了领域知识驱动特征构建的有效性。')

# ===================================================================
# 6. 结果分析
# ===================================================================
add_heading('6. 结果分析', level=1)

add_heading('6.1 SHAP可解释性分析', level=2)
add_para('为理解模型的决策机制，对最佳模型（随机森林）进行SHAP分析。SHAP特征重要性排名如表6所示（完整SHAP可视化见附录experiments/shap_summary.png）。', indent=False)

add_table(
    ['排名', '特征', 'SHAP Importance', '业务含义'],
    [
        ['1', 'tenure_WOE', '0.0592', '在网时长的证据权重编码，反映客户生命周期的非线性影响'],
        ['2', 'RFM_Total', '0.0592', 'RFM综合评分，同时考虑了最近性、频率和消费价值'],
        ['3', 'MonthlyCharges', '0.0557', '月消费金额，高消费客户流失倾向略有增加'],
        ['4', 'PaymentMethod', '0.0557', '支付方式，电子支票支付者流失风险显著偏高'],
        ['5', 'InternetService', '0.0271', '互联网服务类型，光纤用户流失率高于DSL用户'],
        ['6', 'MultipleLines', '0.0271', '多线路服务，与互联网服务类型高度相关'],
        ['7', 'R_Score', '0.0246', '最近性指标，新客户流失风险高于老客户'],
        ['8', 'TotalCharges', '0.0246', '累计消费总额，长期客户的总消费更高但流失率更低'],
    ]
)

add_para('SHAP分析与随机森林特征重要性相互印证，共同指出：tenure（在网时长）、RFM综合评分、MonthlyCharges（月费用）、PaymentMethod（支付方式）和InternetService（互联网服务类型）是驱动客户流失的核心因素。值得注意的是，Contract（合同类型）虽然在相关性分析中排名第一，但在SHAP模型归因中因与其他特征（如tenure、PaymentMethod）存在共线性而权重有所稀释。')

add_heading('6.2 客户聚类画像', level=2)
add_para('K-Means（K = 4）将客户分为四个特征差异显著的群体，画像特征如表7所示（完整可视化见附录experiments/kmeans_clustering.png和cluster_radar.png）。', indent=False)

add_table(
    ['聚类', '人数', '流失率', '月均消费', '在网时长', '特征画像', '策略方向'],
    [
        ['Cluster 0', '2,517 (35.7%)', '30.4%', '$64', '中等', '中等消费、中等忠诚度', '交叉销售增值服务'],
        ['Cluster 1', '2,226 (31.6%)', '21.3%', '$95', '较长', '高消费、高忠诚度、VIP群体', '核心VIP维护'],
        ['Cluster 2', '946 (13.4%)', '60.9%', '$58', '较短', '月付为主、短在网、极高流失风险', '高危优先干预'],
        ['Cluster 3', '1,354 (19.2%)', '4.1%', '$21', '长', '低消费、高稳定性、基础服务用户', '稳定群体保持'],
    ]
)

add_para('聚类结果清晰揭示了客户价值的异质性：Cluster 1是高价值VIP群体（月均$95，流失率21.3%），是运营商利润的核心来源；Cluster 2是急需干预的高危群体（流失率60.9%，946人），若不加干预将造成大量客户流失；Cluster 3是稳定但低价值的群体（流失率仅4.1%）。')

add_heading('6.3 DBSCAN聚类对比', level=2)
add_para('DBSCAN（eps = 0.5, min_samples = 10）的聚类结果与K-Means形成互补（见附录experiments/dbscan_clustering.png）。DBSCAN自动识别了若干个密集客户群体，同时标记出部分"噪声点"——这些客户的行为模式明显异于主流群体，可能对应着异常账户、欺诈行为或特殊需求客户，值得进一步的人工调查。与K-Means的对比表明，两种聚类方法各有优势：K-Means适合业务分组和策略制定，DBSCAN适合异常检测和个案分析。')

add_heading('6.4 错误案例分析', level=2)
add_para('对随机森林模型的错误预测进行抽样分析，发现两类典型的错误模式。假阴性（FN，预测"不流失"实际"流失"）：此类客户往往具有中等在网时长和中等月消费，其行为特征与"忠诚客户"相似但实际已流失。这提示我们需要更多的行为动态特征（如最近三个月的费用变化趋势、服务使用频率变化）来捕捉流失前兆信号。假阳性（FP，预测"流失"实际"不流失"）：此类客户多为月付合同、使用电子支票支付的在网客户。虽然模型将其判定为"高风险"，但这些客户可能因为习惯或惯性而暂时未流失，实际上是"潜在流失客户"，提前干预仍具有预防价值。')

# ===================================================================
# 7. 业务策略建议
# ===================================================================
add_heading('7. 业务策略建议', level=1)
add_para('基于模型分析结果和聚类画像，提出六大方向的业务策略建议，汇总如表8所示。', indent=False)

add_table(
    ['策略', '数据发现', '具体措施', '预期效果'],
    [
        ['合同类型优化', '月付客户流失率远超长期合同客户', '推出有吸引力的1年/2年合同套餐，提供折扣或免费增值服务作为签约激励', '降低流失率15-20%'],
        ['增值服务捆绑', '未订阅OnlineSecurity/TechSupport者流失率显著偏高', '对高危客户免费试用1个月安全服务，体验后引导付费转化', '提升客户粘性和服务切换成本'],
        ['支付方式引导', '电子支票支付者流失率是自动转账者的2-3倍', '引导客户转为自动转账/信用卡支付，给予小额账单折扣作为激励', '提升客户生命周期价值'],
        ['分群精准运营', '不同聚类群体流失风险和价值差异显著', 'Cluster 1: VIP权益升级和专属客服; Cluster 2: 针对性优惠券和主动外呼; Cluster 0: 交叉销售增值服务; Cluster 3: 保持现有服务', '精准投放，提升运营ROI'],
        ['早期预警系统', 'tenure（在网时长）是最重要的预测特征', '建立分级预警：3个月内加强新手引导，6-12月推送个性化优惠，12月以上启动忠诚度计划', '早期干预，降低新客户流失'],
        ['光纤服务优化', 'Fiber Optic用户付费高但流失率也较高', '排查光纤服务体验痛点（网速稳定性、客服响应），针对性改善', '降低高ARPU客户流失率10%'],
    ]
)

add_para('上述建议均基于数据分析和模型输出的客观结论，实施前建议进行小规模A/B测试以验证策略有效性。')

# ===================================================================
# 8. 结论与展望
# ===================================================================
add_heading('8. 结论与展望', level=1)

add_heading('8.1 主要发现', level=2)
add_para('（1）在模型性能方面，随机森林在真实不平衡测试集上以ROC-AUC = 0.8382、Recall = 0.7139的综合表现成为最优模型。SMOTE过采样有效提升了所有模型对流失客户的识别能力，将Recall从基线水平提升了约15至20个百分点。')
add_para('（2）在关键特征方面，SHAP分析和随机森林特征重要性一致表明，tenure（在网时长）、RFM综合评分、MonthlyCharges（月费用）、PaymentMethod（支付方式）和InternetService（互联网服务类型）是驱动客户流失的核心因素。')
add_para('（3）在客户分群方面，K-Means将客户分为4个具有明显差异的群体。其中Cluster 2（946人，流失率60.9%）是需要优先干预的高危群体；Cluster 1（2,226人，月费$95，流失率21.3%）是高价值VIP客户，需要重点维护。')
add_para('（4）在特征工程价值方面，RFM + WOE特征工程为线性模型带来了0.032的显著AUC提升，验证了领域知识驱动特征构建的有效性。')

add_heading('8.2 局限性', level=2)
add_para('（1）数据规模有限：数据集仅包含7,043条记录，对于捕捉所有流失模式可能不够充分。（2）特征维度受限：缺少客户行为时序数据（如通话记录变化、投诉历史），这些动态信息对流失预测可能更为关键。（3）外部因素缺失：未考虑竞争对手活动、经济环境变化等外部因素对客户流失的影响。（4）模型泛化性：实验基于单一数据集，尚未在其他电信企业数据上验证模型的迁移能力。')

add_heading('8.3 未来改进方向', level=2)
add_para('（1）引入MLP、Wide&Deep等神经网络模型进行对比实验，探索深度学习在结构化数据上的潜力。（2）构建客户行为的时间序列特征（如费用的环比/同比变化率、服务使用的增减趋势），捕捉流失的动态前兆。（3）探索Tomek Links、ENN等欠采样方法与SMOTE的组合策略，寻找更优的不平衡处理方案。（4）设计在线A/B测试框架，评估策略的实际业务效果。（5）开发实时流失预测与自动告警系统，实现模型工程化落地。（6）引入外部数据源（区域经济指标、社交媒体情感分析等）丰富特征维度。')

# ===================================================================
# 参考文献
# ===================================================================
add_heading('参考文献', level=1)

refs = [
    '[1] Reichheld F F, Sasser W E. Zero defections: quality comes to services[J]. Harvard Business Review, 1990, 68(5): 105-111.',
    '[2] Lundberg S M, Lee S I. A unified approach to interpreting model predictions[C]. Advances in Neural Information Processing Systems, 2017: 4765-4774.',
    '[3] Lemmens A, Croux C. Bagging and boosting classification trees to predict churn[J]. Journal of Marketing Research, 2006, 43(2): 276-286.',
    '[4] Chen T, Guestrin C. XGBoost: a scalable tree boosting system[C]. Proceedings of the 22nd ACM SIGKDD, 2016: 785-794.',
    '[5] Ke G, Meng Q, Finley T, et al. LightGBM: a highly efficient gradient boosting decision tree[C]. Advances in NIPS, 2017: 3146-3154.',
    '[6] Burez J, Van den Poel D. Handling class imbalance in customer churn prediction[J]. Expert Systems with Applications, 2009, 36(3): 4626-4636.',
    '[7] Chawla N V, Bowyer K W, Hall L O, et al. SMOTE: synthetic minority over-sampling technique[J]. Journal of Artificial Intelligence Research, 2002, 16: 321-357.',
    '[8] Verbeke W, Dejaeger K, Martens D, et al. New insights into churn prediction in the telecommunication sector[J]. European Journal of Operational Research, 2012, 218(1): 211-229.',
    '[9] Wangperawong A, Brun C, Laudy O, et al. Churn analysis using deep convolutional neural networks and autoencoders[J]. arXiv:1604.05377, 2016.',
    '[10] Grinsztajn L, Oyallon E, Varoquaux G. Why do tree-based models still outperform deep learning on typical tabular data?[C]. Advances in NIPS, 2022.',
    '[11] Ribeiro M T, Singh S, Guestrin C. "Why should I trust you?": explaining the predictions of any classifier[C]. Proceedings of the 22nd ACM SIGKDD, 2016: 1135-1144.',
    '[12] IBM. Telco Customer Churn Dataset[DB/OL]. Kaggle, https://www.kaggle.com/datasets/blastchar/telco-customer-churn, 2018.',
    '[13] Breiman L. Random forests[J]. Machine Learning, 2001, 45(1): 5-32.',
    '[14] Pedregosa F, Varoquaux G, Gramfort A, et al. Scikit-learn: machine learning in Python[J]. Journal of Machine Learning Research, 2011, 12: 2825-2830.',
    '[15] McInnes L, Healy J, Melville J. UMAP: uniform manifold approximation and projection for dimension reduction[J]. arXiv:1802.03426, 2018.',
]

for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(ref)
    run.font.size = Pt(10.5)

# ===================================================================
# 附录
# ===================================================================
add_heading('附录', level=1)

add_heading('附录A: 实验环境', level=2)
add_table(
    ['项目', '版本/说明'],
    [
        ['操作系统', 'Windows 11'],
        ['Python', '3.12'],
        ['NumPy', '1.26.4'],
        ['Pandas', '2.2.2'],
        ['Scikit-learn', '1.4.2'],
        ['XGBoost', '3.3.0'],
        ['LightGBM', '4.6.0'],
        ['SHAP', '0.46.0'],
        ['imbalanced-learn', '0.12.3'],
    ]
)

add_heading('附录B: 可复现性说明', level=2)
add_para('所有实验均固定随机种子（random_state = 42），数据按80%/20%分层划分（stratify = y），交叉验证使用5折StratifiedKFold（shuffle = True）。SMOTE仅在训练集上执行，测试集保持原始分布。完整的Notebook和Python脚本见代码附件。运行命令为：pip install -r requirements.txt && jupyter notebook Telco_Customer_Churn_Analysis.ipynb。')

add_heading('附录C: 可视化图表索引', level=2)
charts = [
    'target_distribution.png — 目标变量分布图（count + 饼图）',
    'numeric_distributions.png — 数值特征分布直方图与箱线图',
    'categorical_churn_rate.png — 类别特征流失率对比图（16个子图）',
    'correlation_matrix.png — 特征相关性热力图',
    'feature_importance_preliminary.png — 随机森林特征重要性预分析',
    'feature_importance.png — 最佳模型特征重要性排序',
    'smote_comparison.png — SMOTE过采样前后对比',
    'model_comparison.png — 四种模型性能对比柱状图',
    'roc_curves.png — ROC曲线对比',
    'pr_curves.png — Precision-Recall曲线对比',
    'confusion_matrices.png — 四种模型混淆矩阵',
    'shap_summary.png — SHAP可解释性分析图',
    'elbow_method.png — 肘部法则确定最佳K值',
    'kmeans_clustering.png — K-Means客户聚类散点图 + 饼图',
    'cluster_radar.png — 聚类画像雷达图',
    'dbscan_clustering.png — DBSCAN密度聚类对比图',
]
for c in charts:
    add_bullet(c)

# 保存
doc_path = os.path.join(BASE, '实验报告.docx')
doc.save(doc_path)
print(f'报告已生成: {doc_path}')
